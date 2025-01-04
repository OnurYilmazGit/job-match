import os
import time
import logging
from itertools import islice
from concurrent.futures import ThreadPoolExecutor, as_completed
from fastapi import FastAPI, HTTPException, Request, Depends
from pydantic import BaseModel
from typing import List
import requests
import json
from pdfminer.high_level import extract_text
from dotenv import load_dotenv
from fastapi.middleware.cors import CORSMiddleware
import uuid
import io
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Pt, Inches
from starlette.middleware.sessions import SessionMiddleware

# New imports for Groq
from groq import Groq

# Load environment variables from .env file
load_dotenv()

# Initialize FastAPI app
app = FastAPI()

# Add CORS middleware after creating the FastAPI app
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # Allow requests from the Next.js dev server
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Simulated session storage (in-memory)
session_store = {}

# Function to generate a new secret key
def generate_secret_key():
    return os.urandom(24).hex()

# Function to get session data
def get_session_data(request: Request):
    session_id = request.cookies.get("session_id")
    if session_id and session_id in session_store:
        return session_store[session_id]
    return None

# Function to set session data
def set_session_data(request: Request, data):
    session_id = request.cookies.get("session_id")
    if not session_id:
        session_id = str(uuid.uuid4())
        request.cookies["session_id"] = session_id
    session_store[session_id] = data

# Add session middleware with a dynamically generated secret key
app.add_middleware(SessionMiddleware, secret_key=generate_secret_key())

# Default paths
BASE_DIR = "data"
CV_DIR = os.path.join(BASE_DIR, "sample_cv")
JOBS_DIR = os.path.join(BASE_DIR, "sample_job")

# Ensure directories exist
os.makedirs(CV_DIR, exist_ok=True)
os.makedirs(JOBS_DIR, exist_ok=True)

# Global variables to store data
jobs_data = []
cv_data = {}

# Lightcast API credentials from environment variables
CLIENT_ID = os.getenv("CLIENT_ID")
CLIENT_SECRET = os.getenv("CLIENT_SECRET")
SCOPE = os.getenv("SCOPE")
TOKEN_URL = os.getenv("TOKEN_URL")
SKILLS_URL = os.getenv("SKILLS_URL")

# Global cache dictionary for token pooling
token_pool = []

# Function to generate and populate tokens
def populate_tokens(pool_size=5):
    global token_pool
    token_pool = []
    for _ in range(pool_size):
        payload = {
            "client_id": CLIENT_ID,
            "client_secret": CLIENT_SECRET,
            "grant_type": "client_credentials",
            "scope": SCOPE
        }
        headers = {'Content-Type': 'application/x-www-form-urlencoded'}
        response = requests.post(TOKEN_URL, data=payload, headers=headers)

        if response.ok:
            token = response.json().get("access_token", "")
            token_pool.append(token)
        else:
            logging.error("Failed to generate token")

# Function to extract skills from document with a given token
def extract_skills_from_document_with_token(document_text, token):
    try:
        headers = {'Authorization': f'Bearer {token}', 'Content-Type': 'application/json'}
        data = {"text": document_text}
        response = requests.post(SKILLS_URL, headers=headers, json=data)
        if response.ok:
            return response.json()
        else:
            logging.error(f"Failed to extract skills from document: {response.text}")
            return None
    except Exception as e:
        logging.error(f"Exception in extract_skills_from_document: {e}")
        return None

# Function to extract skill names
def extract_skill_names(api_response):
    return [skill.get('skill', {}).get('name') for skill in api_response.get('data', []) if skill.get('skill', {}).get('name')]

# Function to calculate Jaccard similarity
def jaccard_similarity(set1, set2):
    intersection = len(set1.intersection(set2))
    union = len(set1.union(set2))
    return intersection / union if union else 0

# Function to process job descriptions concurrently
def extract_skills_concurrently(job_descriptions):
    global token_pool
    all_skills = []
    with ThreadPoolExecutor(max_workers=len(token_pool)) as executor:
        future_to_job = {
            executor.submit(extract_skills_from_document_with_token, job.get("description", ""), token_pool[i % len(token_pool)]): job
            for i, job in enumerate(job_descriptions)
        }
        for future in as_completed(future_to_job):
            job = future_to_job[future]
            try:
                response = future.result()
                if response:
                    skills = extract_skill_names(response)
                    all_skills.append((job, skills))
                else:
                    logging.error(f"Failed to extract skills for job: {job.get('positionName')}")
            except Exception as e:
                logging.error(f"Error processing job: {job.get('positionName')} - {e}")
    return all_skills

# Input and output models
class JobMatch(BaseModel):
    id: str
    positionName: str
    company: str
    matchScore: float
    matchedSkills: List[str]
    missingSkills: List[str]
    description: str
    url: str

# Function to get a new access token
def get_access_token():
    payload = {
        "client_id": CLIENT_ID,
        "client_secret": CLIENT_SECRET,
        "grant_type": "client_credentials",
        "scope": SCOPE
    }
    headers = {'Content-Type': 'application/x-www-form-urlencoded'}
    response = requests.post(TOKEN_URL, data=payload, headers=headers)

    if response.ok:
        return response.json().get("access_token", "")
    else:
        logging.error("Failed to generate token")
        return None

def standardize_jobs(jobs):
    standardized_jobs = []
    for job in jobs:
        standardized_jobs.append({
            "id": job.get("job_id", job.get("id")),
            "positionName": job.get("job_title", job.get("positionName")),
            "company": job.get("company_name", job.get("company")),
            "location": job.get("job_location", {}).get("city", job.get("location")),
            "description": job.get("job_description", job.get("description")),
            "url": job.get("job_url", job.get("url")),
            "postedAt": job.get("job_posted_date", job.get("postedAt")),
            "scrapedAt": job.get("scrapedAt"),  # This key doesn't appear in the new JSON, handle as needed.
        })
    return standardized_jobs



@app.post("/upload-jobs/")
async def upload_jobs(batch_size=10):
    global jobs_data
    jobs_data = []
    if not os.path.exists(JOBS_DIR):
        raise HTTPException(status_code=404, detail=f"Directory {JOBS_DIR} not found.")

    file_names = [f for f in os.listdir(JOBS_DIR) if f.endswith('.json')]
    
    def process_file(file_name):
        file_path = os.path.join(JOBS_DIR, file_name)
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                raw_jobs = json.load(file)
                return standardize_jobs(raw_jobs)
        except Exception as e:
            logging.error(f"Error loading or standardizing file {file_name}: {e}")
            return []

    with ThreadPoolExecutor() as executor:
        for i in range(0, len(file_names), batch_size):
            batch = file_names[i:i + batch_size]
            results = executor.map(process_file, batch)
            for standardized_jobs in results:
                jobs_data.extend(standardized_jobs)

    return {"message": f"{len(jobs_data)} jobs uploaded and standardized from {JOBS_DIR}."}


@app.post("/upload-cv/")
async def upload_cv():
    global cv_data
    cv_file = os.path.join(CV_DIR, "CV.pdf")
    if not os.path.exists(cv_file):
        raise HTTPException(status_code=404, detail=f"File {cv_file} not found.")
    
    # Extract text using pdfminer
    content = extract_text(cv_file)
    
    # Print the extracted text for verification
    print("Extracted Text from CV:")
    print(content)
    global cv_info
    cv_info = content
    
    
    token = token_pool[0] if token_pool else get_access_token()
    api_response = extract_skills_from_document_with_token(content, token)
    cv_data["skills"] = extract_skill_names(api_response) if api_response else []
    return {"message": "CV uploaded and skills extracted.", "skills": cv_data["skills"]}

@app.get("/match-jobs/", response_model=List[JobMatch])
async def get_matched_jobs(request: Request, session_data=Depends(get_session_data)):
    if not jobs_data or not cv_data:
        raise HTTPException(status_code=400, detail="Jobs or CV data not uploaded.")

    # Check if matched jobs are already in session
    if session_data and "matched_jobs" in session_data:
        return session_data["matched_jobs"]

    start_time = time.time()

    # Concurrently extract skills from job descriptions
    job_skills_data = extract_skills_concurrently(jobs_data)

    results = []
    for job, job_skills in job_skills_data:
        matched_skills = set(job_skills) & set(cv_data["skills"])
        missing_skills = set(job_skills) - matched_skills
        match_score = int(jaccard_similarity(set(job_skills), set(cv_data["skills"])) * 100)

        # Determine the correct URL to use
        job_url = job.get("externalApplyLink") or job.get("url")

        results.append(JobMatch(
            id=job["id"],
            positionName=job["positionName"],
            company=job["company"],
            matchScore=match_score,
            matchedSkills=list(matched_skills),
            missingSkills=list(missing_skills),
            description=job["description"],
            url=job_url
        ))

    results.sort(key=lambda x: x.matchScore, reverse=True)

    end_time = time.time()
    print(f"Total matching time: {end_time - start_time} seconds")

    # Store matched jobs in session
    if session_data is not None:
        session_data["matched_jobs"] = results
        set_session_data(request, session_data)

    return results

# Store both API keys from environment variables
PRIMARY_GROQ_API_KEY = os.getenv("PRIMARY_GROQ_API_KEY")
SECONDARY_GROQ_API_KEY = os.getenv("SECONDARY_GROQ_API_KEY")

# Function to generate cover letter
@app.post("/generate-cover-letter/")
async def generate_cover_letter(job_id: str):
    # Retrieve the job data based on job_id
    job = next((job for job in jobs_data if job.get("id") == job_id), None)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found.")

    # Ensure CV data is available
    if not cv_data.get("skills"):
        raise HTTPException(status_code=400, detail="CV data not uploaded.")

    # Initialize Groq client with primary API key
    groq_client = Groq(api_key=PRIMARY_GROQ_API_KEY)

    # Prepare the prompt for Groq API
    prompt = f"""
    You are a professional career consultant and expert copywriter. You will craft a concise, powerful, and personalized cover letter for the candidate, Onur YILMAZ, applying for the following role:


    **Job Description:**
    {job}

    **Candidate's CV:**
    {cv_info}
    
    
    **Instructions/Constraints:**
    1. The final cover letter must fit on a single page when rendered in Times New Roman, 12-point font, with normal margins.
    2. Use ONLY the candidate's CV information, the provided job description, and the sample letter below as style inspiration. 
    3. Do NOT request any additional details from the user.
    4. Highlight how the candidate’s experience and skills match the job requirements.
    5. The tone should be professional and enthusiastic without being overly lengthy or repetitive.
    6. Ensure correct grammar, spelling, and punctuation.
    7. Output only the final cover letter text (no placeholders or extra instructions).

    **Sample Letter for Style Reference:**
    Dear Hiring Manager,
    My name is Onur Yılmaz, and I recently completed my Master’s in Informatics at the Technical University of Munich. I was very pleased to see the Cloud Solutions Architect position announced by the Max Planck Computing and Data Facility. I believe that both my academic and professional experiences align directly with the responsibilities and desired qualifications outlined in your job description.

    Academic Background and High-Performance Computing (HPC) Experience
    I completed my master thesis on the SuperMUC (HPC) at TUM, where I developed an ML-based dataset compression method for high-dimensional datasets, achieving a 250-fold reduction in data size. In addition, I successfully completed master-level courses such as Cloud Information Systems, Cloud Computing, and Cloud-Based Data Processing, which have given me a strong foundation in cloud technologies.

    Experience at Intel and Giant Swarm
    Intel (1.5 years): I developed an automation and test monitoring interface that consolidated the results of hundreds of daily CI/CD projects into a single Grafana dashboard, reducing analysis time from one hour to just five minutes. Throughout this process, I gained extensive experience with Docker, Python, Jenkins, Ansible, and Grafana on Linux-based systems.

    Giant Swarm (8-month internship): My work focused on Kubernetes-based cluster management, improving the scalability and reliability of Azure clusters. I also developed automation policies to optimize cloud resource utilization.

    Why MPCDF?
    MPCDF’s OpenStack cloud infrastructure and its projects involving high-performance computing are a perfect match for my thesis research and corporate experience. I am eager to help further develop your existing cloud environment, propose new solutions, and support data-intensive projects across the Max Planck Institutes. In particular, my background in Infrastructure as Code (Terraform, Ansible) and container orchestration (Docker, Kubernetes) closely corresponds to your requirements.

    Personal Motivation and Next Steps
    Beyond my technical background, I have been a jazz drummer for over seven years, an experience that has honed my creativity, collaboration skills, and attention to detail, qualities I bring to every project I undertake. Joining MPCDF would allow me to merge my passion for innovative problem-solving with the opportunity to contribute to research-driven initiatives that support groundbreaking scientific discovery.

    I would be delighted to discuss my experience further and explore how my skills align with MPCDF’s goals. Thank you for considering my application. I look forward to the possibility of contributing to your team and advancing MPCDF’s impactful projects.

    Sincerely,
    Onur Yılmaz


    Now, please draft a final cover letter (one-page maximum) that highlights how the candidate’s background matches this job’s requirements. Use a similar style and structure to the sample letter but tailor it to the specifics of the above job description and the candidate’s CV. 
    """

    try:
        # Perform chat completion using Groq API
        chat_completion = groq_client.chat.completions.create(
            messages=[
                {"role": "user", "content": prompt}
            ],
            model="llama-3.3-70b-specdec",
            max_tokens=500,
            temperature=0.7,
            top_p=0.9,
            frequency_penalty=0.2,
            presence_penalty=0.1
        )

        cover_letter = chat_completion.choices[0].message.content.strip()

        if not cover_letter:
            raise HTTPException(status_code=500, detail="Cover letter generation failed.")

        # Create a DOCX file in-memory
        document = Document()
        
        # Set the default font and size
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Add the cover letter content
        document.add_paragraph(cover_letter)

        # Set the margins
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(1.15)
            section.bottom_margin = Inches(1.15)
            section.left_margin = Inches(1.15)
            section.right_margin = Inches(1.15)

        # Save the document to a BytesIO stream
        docx_stream = io.BytesIO()
        document.save(docx_stream)
        docx_stream.seek(0)

        # Prepare the response headers with the company name in the filename
        company_name = job.get('company', 'Company').replace(' ', '_')
        filename = f"cover_letter_{company_name}.docx"
        headers = {
            'Content-Disposition': f'attachment; filename="{filename}"'
        }

        return StreamingResponse(
            docx_stream,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers=headers
        )

    except Exception as e:
        # Check if the error is a rate limit error
        if "rate_limit_exceeded" in str(e):
            # Switch to secondary API key
            groq_client = Groq(api_key=SECONDARY_GROQ_API_KEY)
            try:
                # Retry the request with the secondary API key
                chat_completion = groq_client.chat.completions.create(
                    messages=[
                        {"role": "user", "content": prompt}
                    ],
                    model="llama-3.3-70b-specdec",
                    max_tokens=500,
                    temperature=0.7,
                    top_p=0.9,
                    frequency_penalty=0.2,
                    presence_penalty=0.1
                )

                cover_letter = chat_completion.choices[0].message.content.strip()

                if not cover_letter:
                    raise HTTPException(status_code=500, detail="Cover letter generation failed.")

                # Create a DOCX file in-memory
                document = Document()
                
                # Set the default font and size
                style = document.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)

                # Add the cover letter content
                document.add_paragraph(cover_letter)

                # Set the margins
                sections = document.sections
                for section in sections:
                    section.top_margin = Inches(1.15)
                    section.bottom_margin = Inches(1.15)
                    section.left_margin = Inches(1.15)
                    section.right_margin = Inches(1.15)

                # Save the document to a BytesIO stream
                docx_stream = io.BytesIO()
                document.save(docx_stream)
                docx_stream.seek(0)

                # Prepare the response headers with the company name in the filename
                company_name = job.get('company', 'Company').replace(' ', '_')
                filename = f"cover_letter_{company_name}.docx"
                headers = {
                    'Content-Disposition': f'attachment; filename="{filename}"'
                }

                return StreamingResponse(
                    docx_stream,
                    media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    headers=headers
                )

            except Exception as e:
                logging.error(f"Groq API Error with secondary key for job_id {job_id}: {e}")
                raise HTTPException(status_code=500, detail="Failed to generate cover letter with secondary key.")
        else:
            logging.error(f"Groq API Error for job_id {job_id}: {e}")
            raise HTTPException(status_code=500, detail="Failed to generate cover letter.")

# Run the app for testing
if __name__ == "__main__":
    populate_tokens(pool_size=2)  # Generate a pool of 2 tokens
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)